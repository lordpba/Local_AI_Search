"""
DocumentLoader — Scans a folder and extracts text from all supported files.
Handles PDF (native + scanned), DOCX, images, TXT, CSV.
"""

import hashlib
import logging
from pathlib import Path
from typing import Generator
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from PIL import Image

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A single extracted document with text and metadata."""
    text: str
    metadata: dict = field(default_factory=dict)

    @property
    def is_empty(self) -> bool:
        return not self.text or len(self.text.strip()) < 10


def file_hash(filepath: Path) -> str:
    """Compute SHA-256 hash of a file for change detection."""
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def scan_folder(folder_path: str, supported_extensions: set[str]) -> list[dict]:
    """
    Scan a folder recursively and return file inventory.
    Returns list of dicts with: path, name, extension, size, hash.
    """
    folder = Path(folder_path)
    if not folder.exists() or not folder.is_dir():
        raise ValueError(f"Cartella non trovata: {folder_path}")

    files = []
    for fp in sorted(folder.rglob("*")):
        if fp.is_file() and fp.suffix.lower() in supported_extensions:
            try:
                files.append({
                    "path": str(fp),
                    "name": fp.name,
                    "extension": fp.suffix.lower(),
                    "size": fp.stat().st_size,
                    "hash": file_hash(fp),
                })
            except (PermissionError, OSError) as e:
                logger.warning(f"Cannot access {fp}: {e}")
    return files


def _classify_file(extension: str) -> str:
    """Classify file type from extension."""
    from app.config import SUPPORTED_EXTENSIONS
    for category, exts in SUPPORTED_EXTENSIONS.items():
        if extension in exts:
            return category
    return "unknown"


def extract_text_from_txt(filepath: Path) -> list[Document]:
    """Extract text from plain text files."""
    try:
        text = filepath.read_text(encoding="utf-8", errors="replace")
        return [Document(
            text=text,
            metadata={"filepath": str(filepath), "filename": filepath.name,
                       "file_type": "text", "page": 1}
        )]
    except Exception as e:
        logger.error(f"Error reading {filepath}: {e}")
        return []


def extract_text_from_csv(filepath: Path) -> list[Document]:
    """Extract text from CSV files — convert rows to readable text."""
    try:
        import pandas as pd
        df = pd.read_csv(filepath, encoding="utf-8", on_bad_lines="skip")
        # Convert each row to a text representation
        lines = []
        columns = list(df.columns)
        lines.append("Colonne: " + ", ".join(columns))
        lines.append("")
        for idx, row in df.iterrows():
            parts = [f"{col}: {val}" for col, val in zip(columns, row) if pd.notna(val)]
            lines.append(f"Riga {idx + 1}: " + " | ".join(parts))

        return [Document(
            text="\n".join(lines),
            metadata={"filepath": str(filepath), "filename": filepath.name,
                       "file_type": "csv", "page": 1, "rows": len(df)}
        )]
    except Exception as e:
        logger.error(f"Error reading CSV {filepath}: {e}")
        return []


def extract_text_from_docx(filepath: Path) -> list[Document]:
    """Extract text from DOCX files."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Also extract tables
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                text += "\n\n" + "\n".join(table_lines)

        return [Document(
            text=text,
            metadata={"filepath": str(filepath), "filename": filepath.name,
                       "file_type": "docx", "page": 1}
        )]
    except Exception as e:
        logger.error(f"Error reading DOCX {filepath}: {e}")
        return []


def extract_text_from_pdf(filepath: Path, min_text_length: int = 50) -> list[Document]:
    """
    Extract text from PDF. Uses native text extraction first.
    Returns pages that need OCR separately (as images).
    """
    documents = []
    pages_needing_ocr = []

    try:
        pdf = fitz.open(str(filepath))
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text", sort=True)

            if text and len(text.strip()) >= min_text_length:
                # Native text extraction succeeded
                documents.append(Document(
                    text=text,
                    metadata={
                        "filepath": str(filepath),
                        "filename": filepath.name,
                        "file_type": "pdf",
                        "page": page_num + 1,
                        "extraction": "native",
                    }
                ))
            else:
                # Page needs OCR — render as image at high DPI for handwriting
                from app.config import PDF_OCR_DPI
                pix = page.get_pixmap(dpi=PDF_OCR_DPI)
                img_data = pix.tobytes("png")
                pages_needing_ocr.append({
                    "page_num": page_num + 1,
                    "image_data": img_data,
                    "filepath": str(filepath),
                    "filename": filepath.name,
                })
        pdf.close()
    except Exception as e:
        logger.error(f"Error reading PDF {filepath}: {e}")

    return documents, pages_needing_ocr


def extract_image_for_ocr(filepath: Path) -> list[dict]:
    """Prepare an image file for OCR processing."""
    try:
        img_data = filepath.read_bytes()
        return [{
            "page_num": 1,
            "image_data": img_data,
            "filepath": str(filepath),
            "filename": filepath.name,
        }]
    except Exception as e:
        logger.error(f"Error reading image {filepath}: {e}")
        return []


def load_documents(
    folder_path: str,
    file_inventory: list[dict],
    progress_callback=None,
) -> tuple[list[Document], list[dict]]:
    """
    Load all documents from the file inventory.

    Returns:
        - documents: list of Document with extracted text
        - ocr_queue: list of dicts with image_data needing OCR
    """
    from app.config import SUPPORTED_EXTENSIONS, PDF_MIN_TEXT_LENGTH

    documents = []
    ocr_queue = []
    total = len(file_inventory)

    for i, file_info in enumerate(file_inventory):
        fp = Path(file_info["path"])
        ext = file_info["extension"]
        category = _classify_file(ext)

        if progress_callback:
            progress_callback(i / total, f"Caricamento {fp.name}...")

        try:
            if category == "text":
                if ext == ".csv":
                    documents.extend(extract_text_from_csv(fp))
                else:
                    documents.extend(extract_text_from_txt(fp))

            elif category == "document":
                if ext == ".pdf":
                    docs, ocr_pages = extract_text_from_pdf(fp, PDF_MIN_TEXT_LENGTH)
                    documents.extend(docs)
                    ocr_queue.extend(ocr_pages)
                elif ext in (".docx", ".doc"):
                    documents.extend(extract_text_from_docx(fp))
                # .odt, .rtf — could be added later

            elif category == "image":
                ocr_queue.extend(extract_image_for_ocr(fp))

        except Exception as e:
            logger.error(f"Failed to process {fp}: {e}")

    if progress_callback:
        progress_callback(1.0, "Caricamento completato")

    return documents, ocr_queue
